# -*- coding: utf-8 -*-
from docx import Document
from pdf2docx import Converter
import requests
from pathlib import Path
import pandas as pd
import os
import re
import arxiv
# doc = Document(r'/home/g19tka13/Desktop/paper/citation/test.docx')
# print(len(doc.paragraphs))
# for i in doc.paragraphs:
#     print(i.text)
# pdf_file = '/home/g19tka13/Desktop/paper/citation/test.pdf'
# docx_file = '/home/g19tka13/Desktop/paper/citation/test2.docx'
# cv = Converter(pdf_file)
# cv.convert(docx_file, start=0, end=None)
# cv.close()
doc_path = Path('/home/g19tka13/Downloads/data/3C/CORE_paper_list')
train_docx_path = os.listdir(doc_path / 'train')
test_docx_path = os.listdir(doc_path / 'test')


def filter_file(f):
    # print(f[-5:])
    # print(f[-5:] == ".docx")
    return True if f[-5:] == '.docx' else False


train_files = list(filter(filter_file, train_docx_path))
test_files = list(filter(filter_file, test_docx_path))
# print(train_files)
# for i in test_files:
#     if str('2132622116') in i:
#         print(i)
# exit()
parent_path = Path('/home/g19tka13/Downloads/data/3C/taskA')
path = parent_path / 'train.csv'
data_train = pd.read_csv(path, sep=',')
data_test = pd.read_csv(str(path).replace('train', 'test'))
data_train_drop = data_train.drop_duplicates(subset='core_id')['core_id'].values.tolist()
data_test_drop = data_test.drop_duplicates(subset='core_id')['core_id'].values.tolist()

# train_files = ['79496059.docx']
'''
    79496059 找不到 引文上下文
'''
def generate_train_data():
    for index, row in data_train.iterrows():
        # print(row['core_id'])
        if row['citation_context'].startswith('#AUTHOR_TAG'):
            query = row['citation_context'].replace('#AUTHOR_TAG', '')
        else:
            query = row['citation_context'].replace('#AUTHOR_TAG', row['cited_author'])
        INCLUDE = False
        for file in train_files:
            paragraph = None
            section_name = []
            if str(row['core_id']) in file:
                INCLUDE = True
                loop_number = 0
                print(file)
                doc = Document(doc_path / 'train/{}'.format(file))
                # doc.settings.odd_and_even_pages_header_footer = False
                # for section in doc.sections:
                #     section.different_first_page_header_footer = False
                #     section.header.is_linked_to_previous = False
                #     section.footer.is_linked_to_previous = False
                for p in doc.paragraphs:
                    # print(p.text)
                    loop_number += 1
                    if paragraph is None:
                        paragraph = p
                    else:
                        if len(p.text) != 0 and 'http' not in p.text and not p.text.isspace():
                            # print(18 * '=', p.text)
                            # if p.text.isupper() or re.match(r'[0-9](\s*)[A-Z]', p.text) :

                            # print(row['core_id'],p.text, p.text.split(' '))
                            split_list = p.text.split(' ')
                            # if row['core_id'] == 141222918 and row['cited_author'] == 'Poggio':
                            #     print(18 * '=', p.text, 18 * '=')
                            # if p.text.isupper() or re.match(r'[0-9]\.(\s*)[A-Z][a-zA-Z]|[0-9](\s*)[A-Z][a-zA-Z]', p.text) or len(split_list) == 1:
                            if p.text.isupper() or re.match(r'[0-9]\.*[0-9]*\s*[A-Z][a-zA-Z]|[0-9]\s*[A-Z][a-zA-Z]|^\s+[A-Z]', p.text) or len(split_list) == 1:
                                # if row['core_id'] == 141222918 and row['cited_author'] == 'Poggio':
                                #     print(18 * '=', p.text)
                                section_name.append(re.sub(r'[0-9]\.\s*|[0-9]\s*', '', p.text))
                            paragraph.add_run(p.text)
                        # print(query[:21])
                        # if index == 7:
                        #     # print(query[:21])
                        #
                        # print('query',query[: 12], 10*'*', p.text)
                        problem_core_id = {141222918: {'Stein':'Multisensory integration in superior', 'Poggio': 'An important challenge, recently highlighted'},
                                           30730382: {'Anderson': 'For example, over the last', 'Qin': 'For example, over the last decade Anderson'},
                                           2188336: {'Ambjorn':'Tensor models have originally been', 'Godfrey':'Tensor models have originally been', 'Sasakura': 'Tensor models have originally been', 'Gustavsson':'3-ary algebras have been', 'Bagger': '3-ary algebras have been'},
                                           5226895: {'Iacono':'High tech AAC systems are expensive'},
                                           81800077: {'Warren':'Observations of active region loops with EIS'},
                                           1385821: {'Wald':'ROOT CAUSE ANALYSIS IN HEALTHCARERCA'}}
                        problem_author = ['Stein', 'Anderson', 'Ambjorn', 'Iacono', 'Warren', 'Poggio','Sasakura', 'Qin', 'Gustavsson','Bagger', 'Godfrey', 'Wald']
                        # problem_core_id = [141222918, 30730382, 2188336]
                        # problem_query = ['Multisensory integration in superior', 'For example, over the last', 'have originally been']
                        if query[:12] in p.text:
                            if (row['cited_author'] == 'Gottlieb' and loop_number == 8) or (row['core_id'] in problem_core_id.keys() and row['cited_author'] in problem_author and row['citation_context'].startswith(problem_core_id[row['core_id']][row['cited_author']]) and problem_core_id[row['core_id']][row['cited_author']] not in p.text):
                                print('citation3', 10 * '=', query[:12])
                                # print(p.text)
                                continue
                            # if index == 99:
                            #     print(index, 10 * '+')
                            # print(section_name)
                            else:
                                section_name = list(filter(lambda q: re.match(r'^[A-Z]+\s*[a-z]*|^[A-Z]+\s*[A-Z]$|^\s+[A-Z]|^[A-Z]([a-zA-Z]+\s*){,10}[A-Za-z]+$', q) and q != '' and (q.isupper() or len(q) <= 50), section_name))
                                print(len(section_name), loop_number)
                                print(row['core_id'])
                                print(row['citation_context'])
                                null_list = [82862948, 82868732, 36208885]
                                if row['core_id'] in null_list or (loop_number <= 5 and len(section_name) == 0):
                                    # if row['core_id'] == 82862948 and (len(section_name) == 0 or section_name[-1] == 'NULL'):
                                    section_name.append('section name NULL')
                                print(section_name)
                                data_train.loc[index, 'section_name'] = section_name[-1]
                                print('section', 10*'+', section_name[-1], index)
                paragraph = re.sub(r'al\.', 'al', paragraph.text)
                sentence_list = re.split(r'[a-zA-Z]+\. ', paragraph)
                for i in range(len(sentence_list)):
                    if query[:12] in sentence_list[i]:
                        if i == 0:
                            data_train.loc[index, 'citation_above'] = 'unknow'
                            data_train.loc[index, 'citation_below'] = sentence_list[i + 1]
                            print('null', sentence_list[i], sentence_list[i + 1])
                            continue
                        if i == len(sentence_list) - 1:
                            data_train.loc[index, 'citation_above'] = sentence_list[i - 1]
                            data_train.loc[index, 'citation_below'] = 'unknow'
                            print(sentence_list[i - 1], sentence_list[i], 'unknow')
                            continue
                        if len(sentence_list) == 1:
                            data_train.loc[index, 'citation_above'] = 'unknow'
                            data_train.loc[index, 'citation_below'] = 'unknow'
                            print('null', sentence_list[i], 'null')
                            continue
                        data_train.loc[index, 'citation_above'] = sentence_list[i - 1]
                        data_train.loc[index, 'citation_below'] = sentence_list[i + 1]
                        print(sentence_list[i - 1], sentence_list[i], sentence_list[i + 1])
                    else:
                        continue
        print(index)
        if INCLUDE is False:
            data_train.loc[index, 'section_name'] = 'unfind'
            data_train.loc[index, 'citation_above'] = 'unabove'
            data_train.loc[index, 'citation_below'] = 'unbelow'

        # print(100*'*')
        # if index == 100:
        #     break
    data_train.to_csv('/home/g19tka13/Downloads/data/3C/taskA/newtrain.csv', sep=',', index=False)


def generate_test_data():
    for index, row in data_test.iterrows():
        if row['citation_context'].startswith('#AUTHOR_TAG'):
            query = row['citation_context'].replace('#AUTHOR_TAG', '')
        else:
            query = row['citation_context'].replace('#AUTHOR_TAG', row['cited_author'])
        INCLUDE = False

        for file in test_files:
            paragraph = None
            section_name = []
            if str(row['core_id']) in file:
                loop_number = 0
                INCLUDE = True
                doc = Document(doc_path / 'test/{}'.format(file))
                for p in doc.paragraphs:
                    loop_number += 1
                    if paragraph is None:
                        paragraph = p
                    else:
                        if len(p.text) != 0 and 'http' not in p.text and not p.text.isspace():
                            split_list = p.text.split(' ')
                            if p.text.isupper() or re.match(r'[0-9]\.*[0-9]*\s*[A-Z][a-zA-Z]|[0-9]\s*[A-Z][a-zA-Z]', p.text) or len(split_list) == 1:
                                section_name.append(re.sub(r'[0-9]\.\s*|[0-9]\s*', '', p.text))
                            paragraph.add_run(p.text)
                        if query[:12] in p.text:
                            if (row['cited_author'] == 'Fox' and loop_number == 22) or (row['cited_author'] == 'Bellucci' and loop_number == 14):  # 在正确的句子之前，query会出现多次，
                                print('citation3', 10 * '=', query[:12])
                                print(p.text)
                                continue
                            print(section_name)
                            # section_name = list(filter(lambda q: re.fullmatch(r'^[A-Z]+\s*[a-z]*|^[A-Z]+\s*[A-Z]$', q) and q != '', section_name))
                            section_name = list(filter(lambda q: re.fullmatch(r'^[A-Z]+\s*[a-zA-Z]*|^[A-Z]+\s*[A-Z]*$|^[A-Z]([a-zA-Z]+\s*){,10}[A-Za-z]+$|^([A-Z]+\s*){,4}[a-zA-Z]*', q) and q != '', section_name))
                            print(len(section_name), loop_number)
                            print(row['core_id'])
                            print(row['citation_context'])
                            null_list = [82105846, 82865416]
                            if (row['core_id'] in null_list and len(section_name)==0) or (loop_number <= 5 and len(section_name) == 0):
                                # if row['core_id'] == 82862948 and (len(section_name) == 0 or section_name[-1] == 'NULL'):
                                section_name.append('section name NULL')
                            print(section_name)
                            data_test.loc[index, 'section_name'] = section_name[-1]
                            print('section', 10 * '+', section_name[-1], index)
                paragraph = re.sub(r'al\.', 'al', paragraph.text)
                sentence_list = re.split(r'[a-zA-Z]+\. ', paragraph)
                for i in range(len(sentence_list)):
                    if query[:12] in sentence_list[i]:
                        if i == len(sentence_list) - 1:
                            data_test.loc[index, 'citation_above'] = sentence_list[i - 1]
                            data_test.loc[index, 'citation_below'] = 'unknow'
                            print(sentence_list[i - 1], sentence_list[i], 'unknow')
                            continue
                        if i == 0:
                            data_test.loc[index, 'citation_above'] = 'unknow'
                            data_test.loc[index, 'citation_below'] = sentence_list[i + 1]
                            print('null', sentence_list[i], sentence_list[i + 1])
                            continue
                        if len(sentence_list) == 1:
                            data_train.loc[index, 'citation_above'] = 'unknow'
                            data_train.loc[index, 'citation_below'] = 'unknow'
                            print('null', sentence_list[i], 'null')
                            continue
                        data_test.loc[index, 'citation_above'] = sentence_list[i - 1]
                        data_test.loc[index, 'citation_below'] = sentence_list[i + 1]
                        print(sentence_list[i - 1], sentence_list[i], sentence_list[i + 1])
                    else:
                        continue
        if INCLUDE is False:
            data_test.loc[index, 'section_name'] = 'unfind'
            data_test.loc[index, 'citation_above'] = 'unabove'
            data_test.loc[index, 'citation_below'] = 'unbelow'
    data_test.to_csv('/home/g19tka13/Downloads/data/3C/taskA/newtest.csv', sep=',', index=False)


if __name__ == "__main__":
    generate_train_data()
'''
    Download paper
'''
# for core_id in data_train_drop:
#     pdf_url = 'https://core.ac.uk:443/api-v2/articles/get/{}/download/pdf?apiKey=RX2IqEY0u9axL7kFOw8bPsetUMcoCByf'.format(core_id)
#     print(pdf_url)
#     r = requests.get(pdf_url)
#     os.mknod('/home/g19tka13/Downloads/data/3C/CORE_paper_list/train_{}.pdf'.format(core_id))
#     with open('/home/g19tka13/Downloads/data/3C/CORE_paper_list/train_{}.pdf'.format(core_id), 'wb') as pdf:
#         pdf.write(r.content)
#
#
# for core_id in data_test_drop:
#     print(core_id)
    # pdf_url = 'https://core.ac.uk:443/api-v2/articles/get/{}/download/pdf?apiKey=RX2IqEY0u9axL7kFOw8bPsetUMcoCByf'.format(core_id)
    # print(pdf_url)
    # r = requests.get(pdf_url)
    # os.mknod('/home/g19tka13/Downloads/data/3C/CORE_paper_list/test_{}.pdf'.format(core_id))
    # with open('/home/g19tka13/Downloads/data/3C/CORE_paper_list/test_{}.pdf'.format(core_id), 'wb') as pdf:
    #     pdf.write(r.content)


# def return_download_from_arxiv():
#     files = os.listdir('/home/g19tka13/Downloads/data/3C/CORE_paper_list/mash')
#     restr = 'train_[0-9]*.pdf'
#     test_list = []
#     train_list = []
#     for train_file in files:
#         pattern = re.compile(r'{}'.format(restr))
#         mash_train = re.findall(pattern, train_file)
#         if len(mash_train) != 0:
#             rep = re.sub(r'_', ".", mash_train[0])
#             train_coreid = rep.split('.')[1]
#             print(train_coreid)
#             train_list.append(train_coreid)
#     for test_file in files:
#         test_str = str(restr).replace('train', 'test')
#         # print(test_str)
#         pattern = re.compile(r'{}'.format(test_str))
#         mash_test = re.findall(pattern, test_file)
#         if len(mash_test) != 0:
#             rep = re.sub(r'_', ".", mash_test[0])
#             test_coreid = rep.split('.')[1]
#             print(test_coreid)
#             test_list.append(test_coreid)
#     # print(len(test_list))
#     return train_list, test_list
#
#
# train_id, test_id = return_download_from_arxiv()
# for id in train_id:
#     data_train_dr = data_train.drop_duplicates(subset='core_id')
#     retrain = data_train_dr.query('core_id=={}'.format(id))
#     # print(retrain['citing_title'].values[-1])
#     # print(retrain['citing_author'].values[-1].split(' ')[-1])
#     query_str = "au:{} AND ti:{}".format(retrain['citing_author'].str.split(' ').values[-1][-1].lower(), retrain['citing_title'].values[-1])
#     query = arxiv.query(query=query_str)
#     if len(query) == 0:
#         print(retrain['citing_title'].values[-1], "train_{}".format(id), retrain['citing_author'].values[-1])
#     # else:
#     #     query = query[0]
#     #     arxiv.download(query, dirpath='/home/g19tka13/Downloads/data/3C/CORE_paper_list',slugify=lambda query: re.sub(r'.*', 'train_{}'.format(id), query.get('title')))
#     #     print(query)
# for id in test_id:
#     data_test_dr = data_test.drop_duplicates(subset='core_id')
#     retrain = data_test_dr.query('core_id=={}'.format(id))
#     # print(retrain['citing_title'].values[-1])
#     # print(retrain['citing_author'].values[-1].split(' ')[-1])
#     query_str = "au:{} AND ti:{}".format(retrain['citing_author'].str.split(' ').values[-1][-1].lower(), retrain['citing_title'].values[-1])
#     query = arxiv.query(query=query_str)
#     if len(query) == 0:
#         print(retrain['citing_title'].values[-1], "test_{}".format(id), retrain['citing_author'].values[-1])
#     # else:
#     #     query = query[0]
#     #     arxiv.download(query, dirpath='/home/g19tka13/Downloads/data/3C/CORE_paper_list',slugify=lambda query: re.sub(r'.*', 'test_{}'.format(id), query.get('title')))

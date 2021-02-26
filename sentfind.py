# -*- coding: utf-8 -*-
from docx import Document
import os
import shutil
from pdf2docx import Converter
import re
from pathlib import Path
import arxiv
import pandas as pd


# query = arxiv.query(query="au:yadav AND ti:Fractal Dimension as a measure of the scale of Homogeneity")
# print(query)
# query = query[0]
# def filename():
#     return 'train_{}'.format(123445678)
# arxiv.download(query, slugify=lambda query: re.sub(r'.*', 'train_{}'.format(123455), query.get('title')))
# pdf_file = '/home/g19tka13/Desktop/paper/citation/test.pdf'
# docx_file = '/home/g19tka13/Downloads/data/3C/CORE_paper_list/test3.docx'
# parse(pdf_file, docx_file, start=0, end=None)
parent_path = Path('/home/g19tka13/Downloads/data/3C/CORE_paper_list/test')
# train = parent_path / 'train.csv'
# files = os.listdir('/home/g19tka13/Downloads/data/3C/CORE_paper_list/mash')
# restr = 'train_[0-9]*.pdf'
# test_list = []
# for train_file in files:
#     pattern = re.compile(r'{}'.format(restr))
#     mash_train = re.findall(pattern, train_file)
#     if len(mash_train) != 0:
#         rep = re.sub(r'_', ".", mash_train[0])
#         coreid = rep.split('.')[1]
#         print(coreid)
# for test_file in files:
#     test_str = str(restr).replace('train', 'test')
#     # print(test_str)
#     pattern = re.compile(r'{}'.format(test_str))
#     mash_test = re.findall(pattern, test_file)
#     if len(mash_test) != 0:
#         rep = re.sub(r'_', ".", mash_test[0])
#         coreid = rep.split('.')[1]
#         print(coreid)
#         test_list.append(coreid)
# print(len(test_list))
files = os.listdir(parent_path)
for file in files:
    # if file[-4:] == ".docx":
        # file = str(file).replace('train_', '')
        # path = parent_path / file
        # fsize = os.path.getsize(path)
        # if fsize <= 1024:
    oldname = parent_path / file
    newname = parent_path / str(file).replace('test_', '')
    os.rename(oldname, newname)
    # shutil.move(str(path).replace('train_', ''), '/home/g19tka13/Downloads/data/3C/CORE_paper_list/train/{}'.format(str(file).replace('train_', '')))

doc = Document(r'/home/g19tka13/Downloads/data/3C/CORE_paper_list/train/81605842.docx')
# str = '/home/g19tka13/Downloads/data/3C/CORE_paper_list/train_1683841.pdf'
# fsize = os.path.getsize(str)
# print(fsize)

# pattern = re.compile('(\.\w*?)tangible learning games.*\\..*?\\.')
a = None
supper = []
doc.settings.odd_and_even_pages_header_footer = False
for section in doc.sections:
    section.different_first_page_header_footer = False
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    # print(section.)
for p in doc.paragraphs:
    split_list = p.text.split(' ')
    if p.text.isupper() or re.match(r'[0-9]\.*[0-9]*\s*[A-Z][a-zA-Z]|[0-9]\s*[A-Z][a-zA-Z]|^\s+[A-Z]', p.text) or len(split_list) == 1:
        print(p.text)
    # if a is None:
    #     a = p
    # else:
    #     if len(p.text) != 0:
    #         a.add_run(p.text)
    # if p.text.isupper():
    #     print(type(p.text))
    #     supper.append(p.text)
    # if 'applied cards to' in p.text:
    #     print(supper)
    #     print(supper[-1])

    # print(p.text)
    # print(p.add_run)
    # print(re.search(pattern, p.text))
    # print(match)
    # print(p.text if p.text.isupper() else len(p.text))
# b = re.sub(r'al.', 'al', a.text)
# list = re.split(r'\.', b)
# print(a.text)
# for i in range(len(list)):
#     if 'applied cards to' in list[i]:
#         print(list[i-1])
#         print(list[i])
#         print(list[i+1])
#         break
#     else:
#         continue

train = pd.read_csv('/home/g19tka13/Downloads/data/3C/taskA/train.csv', sep=',')


